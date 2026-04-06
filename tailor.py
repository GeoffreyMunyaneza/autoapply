"""
tailor.py — Uses Claude API to tailor Geoffrey's resume for a specific job,
then saves a new .docx to the output/resumes folder.

Strategy:
  1. Extract all paragraph text from the base .docx with indices.
  2. Send resume text + job description to Claude.
  3. Claude returns JSON: {paragraph_index: new_text} for changed paragraphs only.
  4. Apply changes to a copy of the original .docx (preserves all formatting).
  5. Save as {Company}_{Title}_{date}.docx
"""

import json
import logging
import os
import re
import shutil
from copy import deepcopy
from datetime import date
from pathlib import Path

import anthropic
from docx import Document

from scraper import Job

logger = logging.getLogger(__name__)

SYSTEM_PROMPT = """You are an expert resume writer helping Geoffrey Munyaneza tailor his resume for a specific job.

Geoffrey's background:
- ML Engineer / AI Engineer with 4+ years production experience
- Also targeting PM roles (uses a separate PM resume)
- Key strengths: computer vision, NLP/LLMs, biometrics/liveness detection (ISO 30107-3 certified), on-device ML
- Amazon Alexa PM intern, Trust Stamp (Nasdaq: IDAI) sole ML engineer, IRCAD Africa
- CMU MS Applied ML, UNC Kenan-Flagler MBA (graduating May 2026)
- Location: Chapel Hill, NC | 3-year OPT starting May 2026

Your task: Given the current resume text and a job description, return a JSON object
that maps paragraph indices (as strings) to replacement text for paragraphs that should change.

Rules:
- NEVER fabricate experience, skills, or credentials Geoffrey does not have.
- Only modify: (1) the professional summary/headline, (2) skills section keywords,
  (3) individual bullet points to inject relevant keywords naturally.
- Keep all quantified achievements (numbers, percentages, metrics) intact.
- Match the writing style and tone of the original resume.
- Do NOT change company names, titles, dates, or contact info.
- Return ONLY a valid JSON object like: {"0": "new text for paragraph 0", "5": "new text for paragraph 5"}
- If no changes are needed, return {}
- Keep changes minimal and high-impact — only change paragraphs where it meaningfully helps ATS/relevance."""


def _extract_paragraphs(doc: Document) -> list[tuple[int, str]]:
    """Return list of (index, text) for non-empty paragraphs."""
    result = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            result.append((i, text))
    return result


def _run_format_key(run) -> tuple:
    """Hashable key capturing a run's inline formatting."""
    try:
        color = str(run.font.color.rgb) if (run.font.color and run.font.color.type) else None
    except Exception:
        color = None
    return (bool(run.bold), bool(run.italic), bool(run.underline), run.font.size, run.font.name, color)


def _apply_changes(doc: Document, changes: dict[int, str]) -> Document:
    """
    Apply text changes to document paragraphs while preserving inline formatting.

    Three strategies depending on run structure:
    1. Single run → simple text replace (no formatting loss possible)
    2. Multiple runs, all same format → collapse into first run (safe)
    3. Multiple runs, mixed format (e.g. bold company name, normal text) →
       distribute new text proportionally across runs by word boundary so
       bold/italic/font changes stay roughly in the same positions.

    Note: para.runs returns new objects on each access, so we snapshot once
    and work with indices to avoid identity-comparison bugs.
    """
    for para_idx, new_text in changes.items():
        if para_idx >= len(doc.paragraphs):
            continue
        para = doc.paragraphs[para_idx]

        # Snapshot runs once — do NOT call para.runs again after this
        all_runs = list(para.runs)
        if not all_runs:
            continue

        # Indices of runs that have actual content
        content_indices = [i for i, r in enumerate(all_runs) if r.text.strip()]
        if not content_indices:
            content_indices = list(range(len(all_runs)))
        if not content_indices:
            continue

        content_runs = [all_runs[i] for i in content_indices]

        # ── Case 1: single content run ──────────────────────────────────────
        if len(content_runs) == 1:
            content_runs[0].text = new_text
            for i, r in enumerate(all_runs):
                if i not in content_indices:
                    r.text = ""
            continue

        # ── Case 2: all content runs share identical formatting ─────────────
        formats = [_run_format_key(r) for r in content_runs]
        if len(set(formats)) == 1:
            content_runs[0].text = new_text
            for r in content_runs[1:]:
                r.text = ""
            for i, r in enumerate(all_runs):
                if i not in content_indices:
                    r.text = ""
            continue

        # ── Case 3: mixed inline formatting — distribute proportionally ─────
        orig_total = sum(len(r.text) for r in content_runs)
        if orig_total == 0:
            content_runs[0].text = new_text
            continue

        proportions = [len(r.text) / orig_total for r in content_runs]
        target_chars = [p * len(new_text) for p in proportions]

        words = new_text.split()
        distributed = [""] * len(content_runs)
        word_cursor = 0

        for i in range(len(content_runs) - 1):
            target = target_chars[i]
            chunk = ""
            while word_cursor < len(words):
                sep = " " if chunk else ""
                candidate = chunk + sep + words[word_cursor]
                # Stop adding words once we exceed target (unless chunk is still empty)
                if len(candidate) > target and chunk:
                    break
                chunk = candidate
                word_cursor += 1
            distributed[i] = chunk

        # Last run gets all remaining words
        distributed[-1] = " ".join(words[word_cursor:])

        for run, text in zip(content_runs, distributed):
            run.text = text

        # Clear non-content runs using the snapshotted index set
        for i, r in enumerate(all_runs):
            if i not in content_indices:
                r.text = ""

    return doc


def tailor_resume(
    job: Job,
    resume_type: str,
    resumes_config: dict,
    output_folder: str,
    claude_config: dict,
    api_key: str,
) -> str | None:
    """
    Tailor the resume for `job`, save to output_folder, return the output path.
    Returns None on failure.
    """
    base_resume_filename = resumes_config.get(resume_type)
    if not base_resume_filename:
        logger.error(f"No resume configured for type '{resume_type}'")
        return None

    base_resume_path = Path(base_resume_filename)
    if not base_resume_path.exists():
        logger.error(f"Base resume not found: {base_resume_path}")
        return None

    # Build output path
    safe_company = re.sub(r'[^\w\s-]', '', job.company).strip().replace(' ', '_')[:30]
    safe_title = re.sub(r'[^\w\s-]', '', job.title).strip().replace(' ', '_')[:30]
    today = date.today().strftime("%Y%m%d")
    output_filename = f"{safe_company}_{safe_title}_{resume_type.upper()}_{today}.docx"
    output_path = Path(output_folder) / output_filename

    output_path.parent.mkdir(parents=True, exist_ok=True)

    # Load the base resume
    doc = Document(str(base_resume_path))
    paragraphs = _extract_paragraphs(doc)

    # Build the resume text for Claude
    resume_text = "\n".join(f"[{idx}] {text}" for idx, text in paragraphs)

    # Truncate job description if very long
    jd = job.description[:4000] if len(job.description) > 4000 else job.description

    user_message = f"""JOB TITLE: {job.title}
COMPANY: {job.company}
LOCATION: {job.location}

JOB DESCRIPTION:
{jd}

---

CURRENT RESUME (paragraph index: text):
{resume_text}

---

Return a JSON object mapping paragraph indices (strings) to new text for only the paragraphs that should change.
Return {{}} if the resume is already well-suited. Do not include any explanation, only the JSON."""

    # Call Claude
    try:
        client = anthropic.Anthropic(api_key=api_key)
        message = client.messages.create(
            model=claude_config.get("model", "claude-haiku-4-5-20251001"),
            max_tokens=claude_config.get("max_tokens", 4096),
            system=SYSTEM_PROMPT,
            messages=[{"role": "user", "content": user_message}],
        )
        response_text = message.content[0].text.strip()
    except Exception as e:
        logger.error(f"Claude API call failed for {job.company} / {job.title}: {e}")
        # Fall back: save unchanged resume
        shutil.copy2(str(base_resume_path), str(output_path))
        logger.info(f"  Saved unmodified resume to {output_path}")
        return str(output_path)

    # Parse JSON response
    try:
        # Strip markdown code fences if Claude wrapped it
        json_text = re.sub(r'^```(?:json)?\s*|\s*```$', '', response_text, flags=re.MULTILINE).strip()
        changes_raw = json.loads(json_text)
        changes = {int(k): v for k, v in changes_raw.items()}
    except (json.JSONDecodeError, ValueError) as e:
        logger.warning(f"Could not parse Claude response as JSON: {e}. Saving unmodified resume.")
        shutil.copy2(str(base_resume_path), str(output_path))
        return str(output_path)

    if not changes:
        # No changes needed — just copy
        save_path = output_path
        for attempt in range(5):
            try:
                shutil.copy2(str(base_resume_path), str(save_path))
                break
            except PermissionError:
                save_path = output_path.with_name(f"{output_path.stem}_{attempt + 1}.docx")
        logger.info(f"  No tailoring needed — saved copy to {save_path.name}")
        return str(save_path)

    # Apply changes to a fresh copy of the document
    doc_copy = Document(str(base_resume_path))
    doc_copy = _apply_changes(doc_copy, changes)

    # If file is locked (e.g. open in Word), append a counter to the name
    save_path = output_path
    for attempt in range(5):
        try:
            doc_copy.save(str(save_path))
            break
        except PermissionError:
            stem = output_path.stem
            save_path = output_path.with_name(f"{stem}_{attempt + 1}.docx")
    else:
        logger.warning(f"  Could not save tailored resume (file locked): {output_filename}")
        return None

    logger.info(f"  Tailored resume saved: {save_path.name} ({len(changes)} changes)")
    return str(save_path)
