"""
tailor.py — Uses Claude API to tailor a resume for a specific job,
then saves a new .docx to output/resumes (auto-approve) or output/pending (review mode).

Strategy:
  1. Extract all paragraph text from the base .docx with indices.
  2. Send resume text + job description to Claude.
  3. Claude returns JSON: {paragraph_index: new_text} for changed paragraphs only.
  4. Apply changes to a copy of the original .docx (preserves all formatting).
  5. Save as {Company}_{Title}_{date}.docx
  6. If auto_approve=False, save to pending/ with a .diff.txt for human review.

User profile (name, background) is read from config.yaml → user_profile section
and passed in at runtime — nothing personal is hardcoded here.
"""

import difflib
import json
import logging
import re
import shutil
from datetime import date
from pathlib import Path

import anthropic
from docx import Document

from core.scraper import Job

# ── Cached Anthropic client (created once per process) ────────────────────────
_client: anthropic.Anthropic | None = None


def _get_client(api_key: str) -> anthropic.Anthropic:
    global _client
    if _client is None:
        _client = anthropic.Anthropic(api_key=api_key)
    return _client

logger = logging.getLogger(__name__)

def _build_resume_prompt(user_profile: dict) -> str:
    """Build the resume tailoring system prompt from the user's profile config."""
    name       = user_profile.get("name", "the applicant")
    background = user_profile.get("background", "").strip()
    bg_section = f"\n{name}'s background:\n{background}\n" if background else ""
    return (
        f"You are an expert resume writer helping {name} tailor their resume for a specific job."
        f"{bg_section}\n"
        "Your task: Given the current resume text and a job description, return a JSON object\n"
        "that maps paragraph indices (as strings) to replacement text for paragraphs that should change.\n\n"
        "Rules:\n"
        f"- NEVER fabricate experience, skills, or credentials {name} does not have.\n"
        "- Only modify: (1) the professional summary/headline, (2) skills section keywords,\n"
        "  (3) individual bullet points to inject relevant keywords naturally.\n"
        "- Keep all quantified achievements (numbers, percentages, metrics) intact.\n"
        "- Match the writing style and tone of the original resume.\n"
        "- Do NOT change company names, titles, dates, or contact info.\n"
        '- Return ONLY a valid JSON object like: {"0": "new text for paragraph 0", "5": "new text"}\n'
        "- If no changes are needed, return {}\n"
        "- Keep changes minimal and high-impact — only change paragraphs where it meaningfully helps."
    )


def _extract_paragraphs(doc: Document) -> list[tuple[int, str]]:
    """Return list of (index, text) for non-empty paragraphs."""
    result = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            result.append((i, text))
    return result


def _run_format_key(run) -> tuple:
    """Hashable key capturing a run's inline formatting."""
    try:
        color = str(run.font.color.rgb) if (run.font.color and run.font.color.type) else None
    except Exception:
        color = None
    return (bool(run.bold), bool(run.italic), bool(run.underline), run.font.size, run.font.name, color)


def _apply_changes(doc: Document, changes: dict[int, str]) -> Document:
    """
    Apply text changes to document paragraphs while preserving inline formatting.

    Three strategies depending on run structure:
    1. Single run → simple text replace (no formatting loss possible)
    2. Multiple runs, all same format → collapse into first run (safe)
    3. Multiple runs, mixed format (e.g. bold company name, normal text) →
       distribute new text proportionally across runs by word boundary so
       bold/italic/font changes stay roughly in the same positions.

    Note: para.runs returns new objects on each access, so we snapshot once
    and work with indices to avoid identity-comparison bugs.
    """
    for para_idx, new_text in changes.items():
        if para_idx >= len(doc.paragraphs):
            continue
        para = doc.paragraphs[para_idx]

        # Snapshot runs once — do NOT call para.runs again after this
        all_runs = list(para.runs)
        if not all_runs:
            continue

        # Indices of runs that have actual content
        content_indices = [i for i, r in enumerate(all_runs) if r.text.strip()]
        if not content_indices:
            content_indices = list(range(len(all_runs)))
        if not content_indices:
            continue

        content_runs = [all_runs[i] for i in content_indices]

        # ── Case 1: single content run ──────────────────────────────────────
        if len(content_runs) == 1:
            content_runs[0].text = new_text
            for i, r in enumerate(all_runs):
                if i not in content_indices:
                    r.text = ""
            continue

        # ── Case 2: all content runs share identical formatting ─────────────
        formats = [_run_format_key(r) for r in content_runs]
        if len(set(formats)) == 1:
            content_runs[0].text = new_text
            for r in content_runs[1:]:
                r.text = ""
            for i, r in enumerate(all_runs):
                if i not in content_indices:
                    r.text = ""
            continue

        # ── Case 3: mixed inline formatting — distribute proportionally ─────
        orig_total = sum(len(r.text) for r in content_runs)
        if orig_total == 0:
            content_runs[0].text = new_text
            continue

        proportions = [len(r.text) / orig_total for r in content_runs]
        target_chars = [p * len(new_text) for p in proportions]

        words = new_text.split()
        distributed = [""] * len(content_runs)
        word_cursor = 0

        for i in range(len(content_runs) - 1):
            target = target_chars[i]
            chunk = ""
            while word_cursor < len(words):
                sep = " " if chunk else ""
                candidate = chunk + sep + words[word_cursor]
                # Stop adding words once we exceed target (unless chunk is still empty)
                if len(candidate) > target and chunk:
                    break
                chunk = candidate
                word_cursor += 1
            distributed[i] = chunk

        # Last run gets all remaining words
        distributed[-1] = " ".join(words[word_cursor:])

        for run, text in zip(content_runs, distributed):
            run.text = text

        # Clear non-content runs using the snapshotted index set
        for i, r in enumerate(all_runs):
            if i not in content_indices:
                r.text = ""

    return doc


def _save_diff(
    original_paragraphs: list[tuple[int, str]],
    changes: dict[int, str],
    diff_path: Path,
    job: "Job",
) -> None:
    """Write a human-readable diff of resume changes to diff_path."""
    orig_map = {idx: text for idx, text in original_paragraphs}
    lines = [
        f"DIFF — {job.title} @ {job.company}",
        f"URL: {job.url}",
        "=" * 70,
        "",
    ]
    for idx, new_text in sorted(changes.items()):
        old_text = orig_map.get(idx, "")
        lines.append(f"[Paragraph {idx}]")
        for dl in difflib.unified_diff(
            old_text.splitlines(),
            new_text.splitlines(),
            fromfile="original",
            tofile="tailored",
            lineterm="",
        ):
            lines.append(dl)
        lines.append("")
    diff_path.write_text("\n".join(lines), encoding="utf-8")


def _build_cover_prompt(user_profile: dict) -> str:
    """Build the cover letter system prompt from the user's profile config."""
    name       = user_profile.get("name", "the applicant")
    background = user_profile.get("background", "").strip()
    bg_section = f"\n{name}'s background:\n{background}\n" if background else ""
    return (
        f"You are helping {name} write a tailored cover letter for a job application."
        f"{bg_section}\n"
        "Write exactly 3 short paragraphs:\n"
        "1. Why THIS specific role and company is exciting (reference the job description directly)\n"
        "2. Two or three most relevant achievements with numbers/metrics\n"
        "3. Enthusiasm + clear call to action\n\n"
        "Rules:\n"
        '- NO generic filler ("I am writing to express...", "I believe I would be a great fit")\n'
        "- Professional but confident tone\n"
        "- Max 200 words total\n"
        "- Return ONLY the body paragraphs — no date, address, salutation, or signature"
    )


def generate_cover_letter(
    job: "Job",
    claude_config: dict,
    api_key: str,
    user_profile: dict | None = None,
) -> str:
    """
    Generate a tailored cover letter for `job` using Claude.
    user_profile: dict from config.yaml → user_profile section.
    Returns the cover letter body text, or empty string on failure.
    """
    if not api_key:
        return ""
    try:
        client = _get_client(api_key)
        message = client.messages.create(
            model=claude_config.get("model", "claude-haiku-4-5-20251001"),
            max_tokens=500,
            system=_build_cover_prompt(user_profile or {}),
            messages=[{
                "role": "user",
                "content": (
                    f"Job: {job.title} at {job.company}\n"
                    f"Location: {job.location}\n\n"
                    f"Job Description:\n{job.description[:3000]}"
                ),
            }],
        )
        text = message.content[0].text.strip()
        logger.info(f"  Cover letter generated ({len(text)} chars)")
        return text
    except Exception as e:
        logger.warning(f"  Cover letter generation failed: {e}")
        return ""


def tailor_resume(
    job: Job,
    resume_type: str,
    resumes_config: dict,
    output_folder: str,
    claude_config: dict,
    api_key: str,
    auto_approve: bool = True,
    pending_folder: str = "",
    generate_cover: bool = False,
    user_profile: dict | None = None,
) -> tuple[str | None, str]:
    """
    Tailor the resume for `job`.

    user_profile: dict from config.yaml → user_profile section (name + background).
    - auto_approve=True  → save directly to output_folder (ready to send)
    - auto_approve=False → save to pending_folder with a .diff.txt for review
    - generate_cover=True → also generate a cover letter with Claude

    Returns (output_path | None, cover_letter_text).
    """
    cover_letter = generate_cover_letter(job, claude_config, api_key, user_profile) if generate_cover else ""

    base_resume_filename = resumes_config.get(resume_type)
    if not base_resume_filename:
        logger.error(f"No resume configured for type '{resume_type}'")
        return None, cover_letter

    base_resume_path = Path(base_resume_filename)
    if not base_resume_path.exists():
        logger.error(f"Base resume not found: {base_resume_path}")
        return None, cover_letter

    # Build output path — pending folder if review required
    safe_company = re.sub(r'[^\w\s-]', '', job.company).strip().replace(' ', '_')[:30]
    safe_title = re.sub(r'[^\w\s-]', '', job.title).strip().replace(' ', '_')[:30]
    today = date.today().strftime("%Y%m%d")
    output_filename = f"{safe_company}_{safe_title}_{resume_type.upper()}_{today}.docx"

    dest_folder = Path(pending_folder) if (not auto_approve and pending_folder) else Path(output_folder)
    dest_folder.mkdir(parents=True, exist_ok=True)
    output_path = dest_folder / output_filename

    # Load the base resume
    doc = Document(str(base_resume_path))
    paragraphs = _extract_paragraphs(doc)

    # Build the resume text for Claude
    resume_text = "\n".join(f"[{idx}] {text}" for idx, text in paragraphs)

    # Truncate job description if very long
    jd = job.description[:4000] if len(job.description) > 4000 else job.description

    user_message = f"""JOB TITLE: {job.title}
COMPANY: {job.company}
LOCATION: {job.location}

JOB DESCRIPTION:
{jd}

---

CURRENT RESUME (paragraph index: text):
{resume_text}

---

Return a JSON object mapping paragraph indices (strings) to new text for only the paragraphs that should change.
Return {{}} if the resume is already well-suited. Do not include any explanation, only the JSON."""

    # Call Claude
    try:
        client = _get_client(api_key)
        message = client.messages.create(
            model=claude_config.get("model", "claude-haiku-4-5-20251001"),
            max_tokens=claude_config.get("max_tokens", 4096),
            system=_build_resume_prompt(user_profile or {}),
            messages=[{"role": "user", "content": user_message}],
        )
        response_text = message.content[0].text.strip()
    except Exception as e:
        logger.error(f"Claude API call failed for {job.company} / {job.title}: {e}")
        # Fall back: save unchanged resume
        shutil.copy2(str(base_resume_path), str(output_path))
        logger.info(f"  Saved unmodified resume to {output_path}")
        return str(output_path), cover_letter

    # Parse JSON response
    try:
        # Strip markdown code fences if Claude wrapped it
        json_text = re.sub(r'^```(?:json)?\s*|\s*```$', '', response_text, flags=re.MULTILINE).strip()
        changes_raw = json.loads(json_text)
        changes = {int(k): v for k, v in changes_raw.items()}
    except (json.JSONDecodeError, ValueError) as e:
        logger.warning(f"Could not parse Claude response as JSON: {e}. Saving unmodified resume.")
        shutil.copy2(str(base_resume_path), str(output_path))
        return str(output_path), cover_letter

    if not changes:
        # No changes needed — just copy to final output (no review needed for unchanged resumes)
        final_path = Path(output_folder) / output_filename
        final_path.parent.mkdir(parents=True, exist_ok=True)
        save_path = final_path
        for attempt in range(5):
            try:
                shutil.copy2(str(base_resume_path), str(save_path))
                break
            except PermissionError:
                save_path = final_path.with_name(f"{final_path.stem}_{attempt + 1}.docx")
        logger.info(f"  No tailoring needed — saved copy to {save_path.name}")
        return str(save_path), cover_letter

    # Apply changes to a fresh copy of the document
    doc_copy = Document(str(base_resume_path))
    doc_copy = _apply_changes(doc_copy, changes)

    # If file is locked (e.g. open in Word), append a counter to the name
    save_path = output_path
    for attempt in range(5):
        try:
            doc_copy.save(str(save_path))
            break
        except PermissionError:
            stem = output_path.stem
            save_path = output_path.with_name(f"{stem}_{attempt + 1}.docx")
    else:
        logger.warning(f"  Could not save tailored resume (file locked): {output_filename}")
        return None, cover_letter

    # Write diff file alongside the resume when pending review
    if not auto_approve and pending_folder:
        diff_path = save_path.with_suffix(".diff.txt")
        _save_diff(paragraphs, changes, diff_path, job)
        logger.info(f"  Pending review → {save_path.name} ({len(changes)} changes) — diff: {diff_path.name}")
    else:
        logger.info(f"  Tailored resume saved: {save_path.name} ({len(changes)} changes)")

    return str(save_path), cover_letter
